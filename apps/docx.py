"Output DOCX for book or item."

import datetime
import io
import json
import struct
import urllib.parse

import docx
import docx
import docx.oxml
import docx.shared
import docx.styles.style
import PIL
import requests
import vl_convert

from fasthtml.common import *

import auth
import books
from books import Book, get_imgs
import components
import constants
from errors import *
import minixml
import users
import utils
from utils import Tx


app, rt = components.get_fast_app()


@rt("/{book:Book}")
def get(request, book: Book):
    "Get the parameters for downloading the book DOCX file."
    auth.authorize(request, *auth.book_view, book=book)

    settings = book.frontmatter.setdefault("docx", {})
    title_page_metadata = bool(settings.get("title_page_metadata", False))
    page_break_level = int(settings.get("page_break_level", 1))
    toc_level = int(settings.get("toc_level", 0))
    toc_level_options = []
    for value in range(0, constants.DOCX_MAX_TOC_LEVEL + 1):
        if value == toc_level:
            toc_level_options.append(Option(str(value), selected=True))
        else:
            toc_level_options.append(Option(str(value)))
    page_break_options = []
    for value in range(0, constants.DOCX_MAX_PAGE_BREAK_LEVEL + 1):
        if value == page_break_level:
            page_break_options.append(Option(str(value), selected=True))
        else:
            page_break_options.append(Option(str(value)))
    footnotes_location = settings.get(
        "footnotes_location", constants.FOOTNOTES_EACH_TEXT
    )
    footnotes_options = []
    for value in constants.FOOTNOTES_LOCATIONS:
        if value == footnotes_location:
            footnotes_options.append(
                Option(Tx(value.capitalize()), value=value, selected=True)
            )
        else:
            footnotes_options.append(Option(Tx(value.capitalize()), value=value))
    reference_font = settings.get("reference_font", constants.NORMAL)
    reference_font_options = []
    for value in constants.FONT_STYLES:
        if value == reference_font:
            reference_font_options.append(
                Option(Tx(value.capitalize()), value=value, selected=True)
            )
        else:
            reference_font_options.append(Option(Tx(value.capitalize()), value=value))
    indexed_font = settings.get("indexed_font", constants.NORMAL)
    indexed_font_options = []
    for value in constants.FONT_STYLES:
        if value == indexed_font:
            indexed_font_options.append(
                Option(Tx(value.capitalize()), value=value, selected=True)
            )
        else:
            indexed_font_options.append(Option(Tx(value.capitalize()), value=value))
    fields = [
        Fieldset(
            Label(
                Input(
                    type="checkbox",
                    name="title_page_metadata",
                    role="switch",
                    checked=title_page_metadata,
                ),
                Tx("Metadata on title page"),
            ),
        ),
        Fieldset(
            Label(Tx("Page break level")),
            Select(*page_break_options, name="page_break_level"),
        ),
        Fieldset(
            Label(Tx("Table of contents level")),
            Select(*toc_level_options, name="toc_level"),
        ),
        Fieldset(
            Label(Tx("Footnotes location")),
            Select(*footnotes_options, name="footnotes_location"),
        ),
        Fieldset(
            Label(Tx("Reference font")),
            Select(*reference_font_options, name="reference_font"),
        ),
        Fieldset(
            Label(Tx("Indexed font")),
            Select(*indexed_font_options, name="indexed_font"),
        ),
    ]

    title = Tx("Book as DOCX file")
    return (
        Title(title),
        components.header(request, title, book=book),
        Main(
            Form(
                *fields,
                components.save_button(Tx("Book as DOCX file")),
                action=f"/docx/{book}",
                method="post",
            ),
            components.cancel_button(f"/book/{book}"),
            cls="container",
        ),
        components.footer(request, book),
    )


@rt("/{book:Book}")
def post(request, book: Book, form: dict):
    "Actually download the book as DOCX file."
    auth.authorize(request, *auth.book_view, book=book)

    settings = book.frontmatter.setdefault("docx", {})
    settings["title_page_metadata"] = bool(form.get("title_page_metadata", False))
    settings["page_break_level"] = int(form.get("page_break_level", 1))
    settings["toc_level"] = int(form.get("toc_level", 0))
    settings["footnotes_location"] = form.get(
        "footnotes_location", constants.FOOTNOTES_EACH_TEXT
    )
    settings["reference_font"] = form.get("reference_font", constants.NORMAL)
    settings["indexed_font"] = form.get("indexed_font", constants.NORMAL)

    # Save settings.
    if auth.authorized(request, *auth.book_edit, book=book):
        book.write()

    return Response(
        content=BookWriter(book).get_content(),
        media_type=constants.DOCX_CONTENT_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{book.title}.docx"'},
    )


@rt("/{book:Book}/{path:path}")
def get(request, book: Book, path: str, position: int = None):
    "Download the item as a DOCX file."
    auth.authorize(request, *auth.book_view, book=book)
    if not path:
        return components.redirect(f"/book/{book}")

    if not "docx" in book.frontmatter:
        raise Error("no DOCX output parameters have been set")

    item = book[path]
    return Response(
        content=ItemWriter(book).get_content(item),
        media_type=constants.DOCX_CONTENT_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{item.title}.docx"'},
    )


class Writer:
    "General DOCX document writer."

    def __init__(self, book):
        self.book = book
        self.references = books.get_refs()

        settings = book.frontmatter.get("docx", {})
        self.title_page_metadata = bool(settings.get("title_page_metadata", False))
        self.page_break_level = int(settings.get("page_break_level", 1))
        self.toc_level = int(settings.get("toc_level", 0))
        self.footnotes_location = settings.get(
            "footnotes_location", constants.FOOTNOTES_EACH_TEXT
        )
        self.reference_font = settings.get("reference_font", constants.NORMAL)
        self.indexed_font = settings.get("indexed_font", constants.NORMAL)
        # Paragraph number
        self.paragraph_number = 0

        # Key: fulltitle; value: dict(label, ast_children)
        self.footnotes = {}
        # 0 = not in footnote; -1 = footnote started; >= 1 = footnote number to start
        self.footnote_def_flag = 0
        # Actually referenced, key: refid; value: reference
        self.referenced = set()
        # Key: canonical; value: dict(id, fulltitle, ordinal)
        self.indexed = {}

        self.document = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        if self.book.language:
            styles_element = self.document.styles.element
            rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
            lang_default = rpr_default.xpath("w:lang")[0]
            lang_default.set(docx.oxml.shared.qn("w:val"), self.book.language)

        # Set to A4 page size.
        section = self.document.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)
        section.footer_distance = docx.shared.Mm(12.7)

        # for cstyle in [
        #     s
        #     for s in self.document.styles
        #     if isinstance(s, docx.styles.style.CharacterStyle)
        #     and not isinstance(s, docx.styles.style._TableStyle)
        # ]:
        #     ic(cstyle, cstyle.name)

        # Modify styles.
        style = self.document.styles["Normal"]
        style.font.name = constants.DOCX_NORMAL_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_NORMAL_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_NORMAL_LINE_SPACING
        )

        # "Body Text": used for TOC entries and for index pages.
        style = self.document.styles["Body Text"]
        style.paragraph_format.space_before = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_BEFORE
        )
        style.paragraph_format.space_after = docx.shared.Pt(
            constants.DOCX_TOC_SPACE_AFTER
        )

        style = self.document.styles["Title"]
        style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        for level in range(1, constants.MAX_LEVEL + 1):
            style = self.document.styles[f"Heading {level}"]
            style.paragraph_format.space_after = docx.shared.Pt(
                2 * (constants.MAX_LEVEL + 1 - level)
            )
            style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        style = self.document.styles["Quote"]
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_QUOTE_INDENT)

        style = self.document.styles["macro"]
        style.font.name = constants.DOCX_CODE_FONT
        style.font.size = docx.shared.Pt(constants.DOCX_CODE_FONT_SIZE)
        style.paragraph_format.line_spacing = docx.shared.Pt(
            constants.DOCX_CODE_LINE_SPACING
        )
        style.paragraph_format.left_indent = docx.shared.Pt(constants.DOCX_CODE_INDENT)

        # Set Dublin core metadata.
        self.document.core_properties.author = ", ".join(self.book.authors)
        self.document.core_properties.created = datetime.datetime.now()
        self.document.core_properties.modified = self.book.modified
        if self.book.language:
            self.document.core_properties.language = self.book.language

        # Display page number in the header.
        # https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.document.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def write_section(self, section, level, skip_page_break=False):
        if section.status == constants.OMITTED:
            return
        if level <= self.page_break_level and not skip_page_break:
            self.document.add_page_break()
        self.write_heading(section.heading, level)
        if section.subtitle:
            self.write_heading(section.subtitle, level + 1)
        self.current_text = section
        self.render_initialize()
        self.render(section.ast)
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            self.write_text_footnotes(section)
        for item in section.items:
            if item.is_section:
                self.write_section(item, level=level + 1)
            else:
                self.write_text(item, level=level + 1)

    def write_text(self, text, level, skip_page_break=False):
        if text.status == constants.OMITTED:
            return
        if level <= self.page_break_level and not skip_page_break:
            self.document.add_page_break()
        if not text.frontmatter.get("suppress_title"):
            self.write_heading(text.heading, level)
            if text.subtitle:
                self.write_heading(text.subtitle, level + 1)
        self.current_text = text
        self.render_initialize()
        self.render(text.ast)
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            self.write_text_footnotes(text)

    def write_heading(self, heading, level):
        level = min(level, constants.MAX_LEVEL)
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        paragraph.add_run(heading)

    def write_text_footnotes(self, text):
        "Footnotes at end of the text."
        assert self.footnotes_location == constants.FOOTNOTES_EACH_TEXT
        try:
            footnotes = self.footnotes[text.fulltitle]
        except KeyError:
            return
        paragraph = self.document.add_heading(
            Tx("Footnotes"), max(3, constants.MAX_LEVEL - 1)
        )
        for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
            self.footnote_def_flag = entry["number"]
            for child in entry["ast_children"]:
                self.render(child)
            self.footnote_def_flag = 0

    def write_chapter_footnotes(self, item):
        "Footnote definitions at the end of a chapter."
        self.footnotes_location == constants.FOOTNOTES_EACH_CHAPTER
        try:
            footnotes = self.footnotes[item.chapter.fulltitle]
        except KeyError:
            return
        self.document.add_page_break()
        self.write_heading(Tx("Footnotes"), 3)
        for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
            self.footnote_def_flag = entry["number"]
            for child in entry["ast_children"]:
                self.render(child)
            self.footnote_def_flag = 0

    def write_book_footnotes(self):
        "Footnote definitions as a separate section at the end of the book."
        assert self.footnotes_location == constants.FOOTNOTES_END_OF_BOOK
        self.document.add_page_break()
        self.write_heading(Tx("Footnotes"), 1)
        for item in self.book.items:
            footnotes = self.footnotes.get(item.fulltitle, {})
            if not footnotes:
                continue
            self.write_heading(item.heading, 2)
            for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
                self.footnote_def_flag = entry["number"]
                for child in entry["ast_children"]:
                    self.render(child)
                self.footnote_def_flag = 0

    def write_references(self):
        self.document.add_page_break()
        self.write_heading(Tx("References"), 1)
        for refid in sorted(self.referenced):
            try:
                reference = self.references[refid]
            except Error:
                continue
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                constants.DOCX_REFERENCE_INDENT
            )
            run = paragraph.add_run(reference["name"])
            run.font.bold = True
            paragraph.add_run("  ")
            self.write_reference_authors(paragraph, reference)
            try:
                method = getattr(self, f"write_reference_{reference['type']}")
            except AttributeError:
                print("unknown", reference["type"])
            else:
                method(paragraph, reference)
            self.write_reference_external_links(paragraph, reference)

    def write_reference_authors(self, paragraph, reference):
        count = len(reference["authors"])
        for pos, author in enumerate(reference["authors"]):
            if pos > 0:
                if pos == count - 1:
                    paragraph.add_run(" & ")
                else:
                    paragraph.add_run(", ")
            paragraph.add_run(utils.short_person_name(author))

    def write_reference_article(self, paragraph, reference):
        paragraph.add_run(" ")
        paragraph.add_run(f"({reference['year']})")
        paragraph.add_run(" ")
        paragraph.add_run(reference.reftitle)
        try:
            run = paragraph.add_run(f"{reference['journal']}")
            run.font.italic = True
            paragraph.add_run(" ")
        except KeyError:
            pass
        try:
            paragraph.add_run(f"{reference['volume']}")
            paragraph.add_run(" ")
        except KeyError:
            pass
        else:
            try:
                paragraph.add_run(f"({reference['number']})")
            except KeyError:
                pass
        try:
            paragraph.add_run(f": pp. {reference['pages'].replace('--', '-')}.")
        except KeyError:
            pass

    def write_reference_book(self, paragraph, reference):
        paragraph.add_run(" ")
        paragraph.add_run(f"({reference['year']})")
        paragraph.add_run(" ")
        run = paragraph.add_run(reference.reftitle)
        run.font.italic = True
        try:
            paragraph.add_run(f" {reference['publisher']}.")
        except KeyError:
            pass
        try:
            paragraph.add_run(f", {reference['edition_published']}")
        except KeyError:
            pass

    def write_reference_link(self, paragraph, reference):
        paragraph.add_run(" ")
        paragraph.add_run(f"({reference['year']})")
        paragraph.add_run(" ")
        run = paragraph.add_run(reference.reftitle)
        run.font.italic = True
        paragraph.add_run(" ")
        try:
            self.add_hyperlink(paragraph, reference["url"], "")
        except KeyError:
            pass
        try:
            paragraph.add_run(f" Accessed {reference['accessed']}.")
        except KeyError:
            pass

    def write_reference_external_links(self, paragraph, reference):
        any_item = False
        if reference.get("url"):
            self.add_hyperlink(paragraph, reference["url"], reference["url"])
            any_item = True
        for key, (label, template) in constants.REFS_LINKS.items():
            try:
                value = reference[key]
                if any_item:
                    paragraph.add_run(", ")
                else:
                    paragraph.add_run(" ")
                self.add_hyperlink(
                    paragraph, template.format(value=value), f"{label}:{value}"
                )
                any_item = True
            except KeyError:
                pass

    def write_indexed(self):
        self.document.add_page_break()
        self.write_heading(Tx("Index"), 1)
        items = sorted(self.indexed.items(), key=lambda i: i[0].casefold())
        for canonical, entries in items:
            paragraph = self.document.add_paragraph(canonical, style="Body Text")
            paragraph.paragraph_format.keep_with_next = True
            entries.sort(key=lambda e: e["ordinal"])
            for entry in entries:
                paragraph = self.document.add_paragraph(
                    entry["heading"], style="Body Text"
                )
                paragraph.paragraph_format.left_indent = docx.shared.Pt(
                    constants.DOCX_INDEXED_INDENT
                )
                if entry is not entries[-1]:
                    paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = docx.shared.Pt(
                constants.DOCX_INDEXED_SPACE_AFTER
            )

    def render_initialize(self):
        "Set up for rendering."
        self.list_stack = []
        self.style_stack = ["Normal"]
        self.bold = False
        self.italic = False
        self.subscript = False
        self.superscript = False

    def render(self, ast):
        "Render the content AST node hierarchy."
        try:
            method = getattr(self, f"render_{ast['element']}")
        except AttributeError:
            print("Could not handle ast", ast)
        else:
            method(ast)

    def render_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.render(child)

    def render_heading(self, ast):
        # XXX Limited implementation; this just handles one child of raw text.
        text = ast["children"][0]["children"]
        self.write_heading(text, ast["level"])

    def render_paragraph(self, ast):
        self.current_paragraph = self.document.add_paragraph()

        if self.footnote_def_flag != 0:
            self.current_paragraph.paragraph_format.left_indent = docx.shared.Pt(
                constants.DOCX_FOOTNOTE_INDENT
            )
            if self.footnote_def_flag > 0:
                self.current_paragraph.paragraph_format.first_line_indent = (
                    -docx.shared.Pt(constants.DOCX_FOOTNOTE_INDENT)
                )
                run = self.current_paragraph.add_run(f"{self.footnote_def_flag}.")
                run.font.bold = True
                self.current_paragraph.add_run(" ")
                self.footnote_def_flag = -1

        if self.list_stack:
            data = self.list_stack[-1]
            levels = min(3, data["levels"])  # Max list levels in predef list styles.
            if data["first_paragraph"]:
                if data["ordered"]:
                    if levels == 1:
                        style = self.document.styles["List Number"]
                    else:
                        style = self.document.styles[f"List Number {levels}"]
                else:
                    if levels == 1:
                        style = self.document.styles["List Bullet"]
                    else:
                        style = self.document.styles[f"List Bullet {levels}"]
            else:
                if levels == 1:
                    style = self.document.styles["List Continue"]
                else:
                    style = self.document.styles[f"List Continue {levels}"]
            data["first_paragraph"] = False
            self.current_paragraph.style = style
        else:
            self.current_paragraph.style = self.style_stack[-1]
        self.paragraph_number += 1
        run = self.current_paragraph.add_run(f"{self.paragraph_number}. ")
        run.style = self.document.styles["Intense Quote Char"]
        for child in ast["children"]:
            self.render(child)

    def render_raw_text(self, ast):
        line = ast["children"]
        line = line.rstrip("\n")
        run = self.current_paragraph.add_run(line)
        if self.bold:
            run.font.bold = True
        if self.italic:
            run.font.italic = True
        if self.subscript:
            run.font.subscript = True
        if self.superscript:
            run.font.superscript = True

    def render_blank_line(self, ast):
        pass

    def render_quote(self, ast):
        self.style_stack.append("Quote")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_code_span(self, ast):
        run = self.current_paragraph.add_run(ast["children"])
        run.style = self.document.styles["Macro Text Char"]

    def render_code_block(self, ast):
        self.current_paragraph = self.document.add_paragraph(style="macro")
        self.style_stack.append("macro")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_fenced_code(self, ast):
        self.current_paragraph = self.document.add_paragraph(style="macro")
        self.style_stack.append("macro")
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_image(self, ast):
        try:
            # Fetch image from the web.
            if urllib.parse.urlparse(ast["dest"]).scheme:
                response = requests.get(ast["dest"])
                if response.status_code != HTTP.OK:
                    raise ValueError(f"Could not fetch image '{ast['dest']}'")
                if response.headers["Content-Type"] not in (
                    constants.PNG_CONTENT_TYPE,
                    constants.JPEG_CONTENT_TYPE,
                ):
                    raise ValueError(
                        f"Cannot handle image '{ast['dest']}' with content type '{response.headers['Content-Type']}'"
                    )
                self.add_image(response.content, img["docx"]["scale_factor"], ast)

            # Use image from the image library.
            elif ast["dest"] in get_imgs():
                img = get_imgs()[ast["dest"]]
                scale_factor = img["docx"]["scale_factor"]

                if img["content_type"] in (
                    constants.SVG_CONTENT_TYPE,
                    constants.JSON_CONTENT_TYPE,
                ):
                    # SVG image.
                    if img["content_type"] == constants.SVG_CONTENT_TYPE:
                        # SVG in image library has already been checked for validity.
                        root = minixml.parse_content(img["data"])

                    # Vega-Lite plot.
                    else:
                        # JSON in image library has already been checked for validity.
                        vl_spec = json.loads(img["data"])
                        root = minixml.parse_content(
                            vl_convert.vegalite_to_svg(vl_spec)
                        )

                    # Set viewbox so that scaling behaves.
                    root["viewBox"] = f"0 0 {root['width']} {root['height']}"

                    # Scale width and height in SVG element.
                    rendering_factor = img["docx"]["png_rendering_factor"]
                    root["width"] = rendering_factor * float(root["width"])
                    root["height"] = rendering_factor * float(root["height"])

                    self.add_image(
                        vl_convert.svg_to_png(repr(root)),
                        ast,
                        scale_factor / rendering_factor,
                    )

                # JPEG or PNG.
                elif img["content_type"] in (
                    constants.PNG_CONTENT_TYPE,
                    constants.JPEG_CONTENT_TYPE,
                ):
                    self.add_image(
                        base64.standard_b64decode(img["data"]), ast, scale_factor
                    )
                else:
                    raise ValueError(f"Cannot handle image {img['content_type']}")
            else:
                raise ValueError(f"No such image '{ast['dest']}'")

        except ValueError as error:
            self.current_paragraph = self.document.add_paragraph(style="macro")
            self.current_paragraph.add_run(str(error))

    def add_image(self, image_data, ast, factor):
        image = io.BytesIO(image_data)
        width, height = PIL.Image.open(image).size
        width = docx.shared.Pt(factor * width)
        height = docx.shared.Pt(factor * height)
        paragraph = self.document.add_paragraph()
        # This is a kludge; seems required to avoid an obscure 'docx' bug?
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(image, width=width, height=height)
        if ast["children"]:
            paragraph.paragraph_format.keep_with_next = True
            self.current_paragraph = self.document.add_paragraph(style="Normal")
            for child in ast["children"]:
                self.render(child)

    def render_emphasis(self, ast):
        self.italic = True
        for child in ast["children"]:
            self.render(child)
        self.italic = False

    def render_strong_emphasis(self, ast):
        self.bold = True
        for child in ast["children"]:
            self.render(child)
        self.bold = False

    def render_subscript(self, ast):
        self.subscript = True
        for child in ast["children"]:
            self.render(child)
        self.subscript = False

    def render_superscript(self, ast):
        self.superscript = True
        for child in ast["children"]:
            self.render(child)
        self.superscript = False

    def render_emdash(self, ast):
        self.current_paragraph.add_run(constants.EM_DASH)

    def render_line_break(self, ast):
        if ast.get("soft"):
            self.current_paragraph.add_run(" ")
        else:
            self.current_paragraph.add_run("\n")

    def render_thematic_break(self, ast):
        paragraph = self.document.add_paragraph(constants.EM_DASH * 20)
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def render_link(self, ast):
        # This handles only raw text within a link, nothing else.
        raw_text = []
        for child in ast["children"]:
            if child["element"] == "raw_text":
                raw_text.append(child["children"])
        self.add_hyperlink(self.current_paragraph, ast["dest"], "".join(raw_text))

    def render_list(self, ast):
        data = dict(
            ordered=ast["ordered"],
            bullet=ast["bullet"],  # XXX Currently not used.
            start=ast["start"],  # XXX Currently not used.
            tight=ast["tight"],  # XXX Currently not used.
            count=0,  # XXX Currently not used.
            levels=len(self.list_stack) + 1,
        )
        self.list_stack.append(data)
        for child in ast["children"]:
            self.render(child)
        self.list_stack.pop()

    def render_list_item(self, ast):
        data = self.list_stack[-1]
        data["count"] += 1  # Currently useless.
        data["first_paragraph"] = True
        for child in ast["children"]:
            self.render(child)

    def render_indexed(self, ast):
        entries = self.indexed.setdefault(ast["canonical"], [])
        entries.append(
            dict(
                ordinal=self.current_text.ordinal,
                fulltitle=self.current_text.fulltitle,
                heading=self.current_text.heading,
            )
        )
        run = self.current_paragraph.add_run(ast["term"])
        if self.indexed_font == constants.ITALIC:
            run.font.italic = True
        elif self.indexed_font == constants.BOLD:
            run.font.bold = True
        elif self.indexed_font == constants.UNDERLINE:
            run.font.underline = True

    def render_footnote_ref(self, ast):
        # The label is used only for lookup; number is used for output.
        label = ast["label"]
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            entries = self.footnotes.setdefault(self.current_text.fulltitle, {})
            number = len(entries) + 1
            key = label
        elif self.footnotes_location in (
            constants.FOOTNOTES_EACH_CHAPTER,
            constants.FOOTNOTES_END_OF_BOOK,
        ):
            fulltitle = self.current_text.chapter.fulltitle
            entries = self.footnotes.setdefault(fulltitle, {})
            number = len(entries) + 1
            key = f"{fulltitle}-{label}"
        entries[key] = dict(label=label, number=number)
        run = self.current_paragraph.add_run(str(number))
        run.font.superscript = True
        run.font.bold = True

    def render_footnote_def(self, ast):
        label = ast["label"]
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            fulltitle = self.current_text.fulltitle
            key = label
        elif self.footnotes_location in (
            constants.FOOTNOTES_EACH_CHAPTER,
            constants.FOOTNOTES_END_OF_BOOK,
        ):
            fulltitle = self.current_text.chapter.fulltitle
            key = f"{fulltitle}-{label}"
        # Footnote def may be missing.
        try:
            self.footnotes[fulltitle][key]["ast_children"] = ast["children"]
        except KeyError:
            pass

    def render_reference(self, ast):
        if ast["id"] in self.references:
            self.referenced.add(ast["id"])
            run = self.current_paragraph.add_run(ast["name"])
            if self.reference_font == constants.ITALIC:
                run.font.italic = True
            elif self.reference_font == constants.BOLD:
                run.font.bold = True
            elif self.reference_font == constants.UNDERLINE:
                run.font.underline = True
        else:
            self.current_paragraph.add_run(f'??? no such refid {ast["name"]} ???')

    # https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
    def add_hyperlink(self, paragraph, url, text, color="2222FF", underline=True):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value.
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        # Create the w:hyperlink tag and add needed values.
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(
            docx.oxml.shared.qn("r:id"),
            r_id,
        )

        # Create a w:r element.
        new_run = docx.oxml.shared.OxmlElement("w:r")

        # Create a new w:rPr element.
        rPr = docx.oxml.shared.OxmlElement("w:rPr")

        # Add color if it is given.
        if not color is None:
            c = docx.oxml.shared.OxmlElement("w:color")
            c.set(docx.oxml.shared.qn("w:val"), color)
            rPr.append(c)

        # Remove underlining if it is requested.
        # XXX Does not seem to work? /Per Kraulis
        if not underline:
            u = docx.oxml.shared.OxmlElement("w:u")
            u.set(docx.oxml.shared.qn("w:val"), "none")
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element.
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink


class BookWriter(Writer):
    "DOCX book writer."

    def get_content(self):
        "Create the DOCX document of the book return its content."
        paragraph = self.document.add_paragraph(style="Title")
        run = paragraph.add_run(self.book.title)

        if self.book.subtitle:
            paragraph = self.document.add_paragraph(style="Heading 1")
            paragraph.add_run(self.book.subtitle)

        # Split authors into runs to allow line break between them.
        paragraph = self.document.add_paragraph(style="Heading 2")
        for author in self.book.authors:
            paragraph.add_run(author)
            if author != self.book.authors[-1]:
                paragraph.add_run(", ")

        self.render_initialize()
        self.render(self.book.ast)

        if self.title_page_metadata:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_before = docx.shared.Pt(
                constants.DOCX_METADATA_SPACER
            )

            run = paragraph.add_run(f"{Tx('Status')}: {Tx(self.book.status)}")
            run.font.italic = True
            run.add_break()
            run = paragraph.add_run(f"{Tx('Created')}: {utils.str_datetime_display()}")
            run.font.italic = True
            run.add_break()
            run = paragraph.add_run(
                f'{Tx("Modified")}: {utils.str_datetime_display(self.book.modified)}'
            )
            run.font.italic = True

        # Write table of contents (TOC) page(s).
        if self.toc_level:
            self.document.add_page_break()
            self.write_heading(Tx("Contents"), 1)
            for item in self.book:
                if item.level > self.toc_level:
                    continue
                if item.status == constants.OMITTED:
                    continue
                paragraph = self.document.add_paragraph(style="Body Text")
                paragraph.paragraph_format.left_indent = docx.shared.Pt(
                    constants.DOCX_TOC_INDENT * item.level
                )
                paragraph.paragraph_format.first_line_indent = -docx.shared.Pt(
                    constants.DOCX_TOC_INDENT
                )
                paragraph.add_run(item.heading)
            # At this stage it is not known if any references or indexed.
            self.document.add_paragraph(Tx("References"), style="Body Text")
            self.document.add_paragraph(Tx("Index"), style="Body Text")

        # First-level items are chapters.
        for item in self.book.items:
            if item.status == constants.OMITTED:
                continue

            if item.is_section:
                self.write_section(item, level=item.level)
            else:
                self.write_text(item, level=item.level)

            if self.footnotes_location == constants.FOOTNOTES_EACH_CHAPTER:
                self.write_chapter_footnotes(item)

        if self.footnotes_location == constants.FOOTNOTES_END_OF_BOOK:
            self.write_book_footnotes()

        # References pages written, even if empty, since TOC pages contains item.
        if self.toc_level or self.referenced:
            self.write_references()
        # Indexed pages written, even if empty, since TOC pages contains item.
        if self.toc_level or self.indexed:
            self.write_indexed()

        output = io.BytesIO()
        self.document.save(output)
        return output.getvalue()


class ItemWriter(Writer):
    "DOCX item (section or text) writer."

    def get_content(self, item):
        "Create the DOCX document of the given item and return its content."
        # Change to the modified datetime of the item
        self.document.core_properties.modified = item.modified

        # Force footnotes at end of each text.
        self.footnotes_location = constants.FOOTNOTES_EACH_TEXT

        if item.is_section:
            self.write_section(item, level=item.level, skip_page_break=True)
        else:
            self.write_text(item, level=item.level, skip_page_break=True)

        # Here it is possible to skip references pages, in none.
        if self.referenced:
            self.write_references()
        # Here it is possible to skip indexed pages, in none.
        if self.indexed:
            self.write_indexed()

        output = io.BytesIO()
        self.document.save(output)
        return output.getvalue()
